"""Text extraction utilities for various file formats."""

import os
import logging
import tempfile
from pathlib import Path
from typing import Optional
import aiofiles

logger = logging.getLogger(__name__)


def is_url(path: str) -> bool:
    """Check if path is a URL."""
    return path.startswith("http://") or path.startswith("https://")


async def extract_text_from_bytes(content: bytes, file_type: str) -> str:
    """
    Extract text content directly from file bytes.
    
    Args:
        content: File content as bytes
        file_type: File extension (pdf, txt, docx, md)
    
    Returns:
        Extracted text content
    """
    file_type = file_type.lower()
    temp_file = None
    
    try:
        # Write bytes to temp file for processing
        import tempfile
        suffix = f".{file_type}"
        with tempfile.NamedTemporaryFile(mode='wb', suffix=suffix, delete=False) as f:
            f.write(content)
            temp_file = f.name
        
        logger.info(f"Processing file from memory ({len(content)} bytes)")
        
        if file_type == "pdf":
            return await extract_from_pdf(temp_file)
        elif file_type in ["txt", "md"]:
            return await extract_from_text(temp_file)
        elif file_type in ["doc", "docx"]:
            return await extract_from_docx(temp_file)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    finally:
        # Clean up temp file
        if temp_file and os.path.exists(temp_file):
            try:
                os.remove(temp_file)
                logger.info(f"Cleaned up temp file: {temp_file}")
            except Exception as e:
                logger.warning(f"Failed to clean up temp file: {e}")


async def extract_text_from_file(file_path: str, file_type: str) -> str:
    """
    Extract text content from a file based on its type.
    This is kept for backward compatibility when processing from URLs.
    
    Args:
        file_path: Path to the file or Supabase URL
        file_type: File extension (pdf, txt, docx, md)
    
    Returns:
        Extracted text content
    """
    file_type = file_type.lower()
    temp_file = None
    
    try:
        # If it's a URL, download to temp file first
        if is_url(file_path):
            from app.services.storage_service import storage_service
            suffix = f".{file_type}"
            temp_file = await storage_service.download_to_temp_file(file_path, suffix=suffix)
            file_path = temp_file
            logger.info(f"Downloaded file to temp: {temp_file}")
        
        if file_type == "pdf":
            return await extract_from_pdf(file_path)
        elif file_type in ["txt", "md"]:
            return await extract_from_text(file_path)
        elif file_type in ["doc", "docx"]:
            return await extract_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    finally:
        # Clean up temp file
        if temp_file and os.path.exists(temp_file):
            try:
                os.remove(temp_file)
                logger.info(f"Cleaned up temp file: {temp_file}")
            except Exception as e:
                logger.warning(f"Failed to clean up temp file: {e}")


async def extract_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file."""
    try:
        import PyPDF2
        
        text_parts = []
        
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
        
        full_text = "\n\n".join(text_parts)
        logger.info(f"Extracted {len(full_text)} characters from PDF: {file_path}")
        return full_text
        
    except ImportError:
        logger.error("PyPDF2 not installed. Install with: pip install PyPDF2")
        raise
    except Exception as e:
        logger.error(f"Failed to extract text from PDF {file_path}: {e}")
        raise


async def extract_from_text(file_path: str) -> str:
    """Extract text from a plain text or markdown file."""
    try:
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
            content = await file.read()
        
        logger.info(f"Extracted {len(content)} characters from text file: {file_path}")
        return content
        
    except UnicodeDecodeError:
        # Try with different encoding
        async with aiofiles.open(file_path, 'r', encoding='latin-1') as file:
            content = await file.read()
        return content
    except Exception as e:
        logger.error(f"Failed to extract text from file {file_path}: {e}")
        raise


async def extract_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    try:
        from docx import Document
        
        doc = Document(file_path)
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        full_text = "\n\n".join(text_parts)
        logger.info(f"Extracted {len(full_text)} characters from DOCX: {file_path}")
        return full_text
        
    except ImportError:
        logger.error("python-docx not installed. Install with: pip install python-docx")
        raise
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX {file_path}: {e}")
        raise


def get_page_number_from_text(text: str) -> Optional[int]:
    """Extract page number from text if it contains a page marker."""
    import re
    match = re.search(r'\[Page (\d+)\]', text)
    if match:
        return int(match.group(1))
    return None
